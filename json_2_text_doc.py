# -*- coding: utf-8 -*-
"""
Created on Mon Sep 30 13:43:13 2019

@author: 726094
"""

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import re
import random
import json
import os

    
# In[Converting json into Citation Output text file]:
def getjsontotext():
    """
    <CITATION-INPUT_FILE_NAME-CITE_ID> :=
    	TYPE: citation
    	SUBTYPE:
    	ANAPHORIC:
    	CO-REF:
    	NAME: "cite string"
        
        <CITATION-ALAAP-4MG8-5M80-0039-4487-00000-00-8070> :=
    	TYPE: citation
    	SUBTYPE: judicialcourtdecision
    	ANAPHORIC: false
    	CO-REF: NA
    	NAME: "Floyd v. Title Exchange & Pawn of Anniston, Inc., 620 So. 2d 576 (Ala. 1993)" ##14747#14827#
    """
    
    #Reading the json file
    json_data = 'outputs/result_citation_data_coref_out.json'
    
    # read file
    with open(json_data, 'r') as data:
        jdata=data.read()
    # parse file
    obj = json.loads(jdata)
    uniqueNames = [];
    for iv in obj['citation']:
        if(iv["filename"] not in uniqueNames):
             uniqueNames.append(iv["filename"]);
  
        
    text_data = []
    for i in range(len(obj['citation'])):
        filename = obj['citation'][i]['filename']
        filename = filename.replace('.txt', '')
        cite_id = obj['citation'][i]['citation_id']
        ctype = obj['citation'][i]['subtype']
        anaphoric = obj['citation'][i]['anaphoric']['status']
        co_ref = obj['citation'][i]['anaphoric']['co-ref']
        name = obj['citation'][i]['name']
        start = (obj['citation'][i]['citation_textstartid'])
        length = (obj['citation'][i]['citation_textlength'])
        
        data = ('<CITATION-'+filename+'-'+str(cite_id)+'> := \n'
        '\t TYPE: citation \n'
        '\t SUBTYPE: '+ctype+'\n'
        '\t ANAPHORIC: '+anaphoric+'\n'
        '\t CO-REF: '+str(co_ref)+'\n'
        '\t NAME: '+name+' ##'+str(start)+'#'+str(start+length)+'# \n')
        text_data.append(data)
    #    print(data)
    filename = filename.replace('.xml.xml','')
    f = open("outputs/"+filename+"-text.txt", "w")
    f.writelines(text_data)
    f.close()
    print('Text Document created: '+filename+'-text.txt')

# In[Converting text to word and highlight using json info]:

def getjsontodocx():
    
    #Reading the json file
    json_data = 'outputs/result_citation_data_coref_out.json'
    input_path = 'input/'
    # read file
    with open(json_data, 'r') as data:
        jdata=data.read()
    # parse file
    obj = json.loads(jdata)
    uniqueNames = [];
    for iv in obj['citation']:
        if(iv["filename"] not in uniqueNames):
             uniqueNames.append(iv["filename"]);
    
    raw_text_file = uniqueNames[0]
        
    #Reading the raw text file
    fname = raw_text_file.replace('xml.xml','txt')
    myfile = open(input_path + fname).read()
#    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
#    print(myfile)
    colors = [
    WD_COLOR_INDEX.GREEN,
    WD_COLOR_INDEX.PINK,
    WD_COLOR_INDEX.TEAL,
    WD_COLOR_INDEX.TURQUOISE,
    WD_COLOR_INDEX.YELLOW]
    
    citations = []
    highlight_color = []
    
    for i in range(len(obj['citation'])):
#        print(raw_text_file)
        if(obj['citation'][i]['filename'] == raw_text_file):
            rand = int(random.uniform(0,4))
            js_object = {
                    'name': obj['citation'][i]['name'],
                    'anaphoric_stat': obj['citation'][i]['anaphoric']['status'],
                    'co_ref': obj['citation'][i]['anaphoric']['co-ref'],
                    'cite_id': obj['citation'][i]['citation_id']  
                    }

            if(js_object['anaphoric_stat'] == 'T' and js_object['co_ref'] != 'nan'):
                cite_id = js_object['co_ref']
                highlight_color.append(highlight_color[citations.index(js_object['cite_id'] == cite_id)])
                
            elif(js_object in citations):
                highlight_color.append(highlight_color[citations.index(js_object)])
            else:
                i = 0
                while (i<50):
                    i+=1
                    col = int(random.uniform(0,4))
                    if(rand not in highlight_color):
                        break
                highlight_color.append(colors[col])            
            citations.append(js_object)

    document = Document()
    

        
    # Create a new paragraph with "helloworld" highlighted
    p2 = document.add_paragraph()
    substrings = []
            
    startid = 0
    length = 0
    for i in range(len(obj['citation'])):
        startid = (obj['citation'][i]['citation_textstartid'])
        length = (obj['citation'][i]['citation_textlength'])
        if i==0:
            substrings.append(myfile[0 : (startid)])
        else:
            new_id = obj['citation'][i-1]['citation_textstartid'] + obj['citation'][i-1]['citation_textlength']
            substrings.append(myfile[new_id : (startid - 1)])
        
    last_id = (startid + length + 2)
    substrings.append(myfile[last_id : len(myfile)])
    #print(substrings)
    
    for i in range(len(citations)):
        p2.add_run(' '+substrings[i]+' ')
        font = p2.add_run(obj['citation'][i]['name']).font
    #    font.highlight_color = WD_COLOR_INDEX.RED
        font.highlight_color = highlight_color[i]
    p2.add_run(substrings[-1])
    
    fname = fname.replace('.txt','')
    document.save('outputs/'+fname+'-word.docx')
    print('Word Document created: '+fname+'-word.docx')


























